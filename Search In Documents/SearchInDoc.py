pip install docx2txt

pip install python-docx

import docx2txt

import docx


from docx.enum.text import WD_COLOR_INDEX
from docx import Document

text = input("Please enter text: ")
import os
path = '/content/sample_data/documents'  //Put the path to the files here
def searchText(path):
    
    os.chdir(path)
    files = os.listdir()
    #print(files)
    for file_name in files:
        #print(file_name)
        abs_path = os.path.abspath(file_name)
        
        if os.path.isdir(abs_path):
            searchText(abs_path)
            
        if os.path.isfile(abs_path):     
            doc = Document(file_name)
           
            for paragraph in doc.paragraphs:
                if text in paragraph.text:
                    print(text + " word found in this path " + file_name)
                    x = paragraph.text.split(text)
                    paragraph.clear()
                    for i in range(len(x)-1):
                        paragraph.add_run(x[i])
                        paragraph.add_run(" "+text+" ") .font.highlight_color = WD_COLOR_INDEX.YELLOW
                        paragraph.add_run(x[i+1])
                        doc.save("/content/sample_data/output/"+file_name)  //Put the output file path here            
                                       

    pass
searchText(path)
